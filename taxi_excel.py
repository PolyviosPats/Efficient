import docx
import os
import openpyxl

l =[]
doc = ''
m = 3
p = 4

wb = openpyxl.load_workbook('kod_taxis.xlsx')

sheet = wb["Sheet1"]

os.chdir('C:\Users\Polis\Desktop\Documents\kodtaxis')

for file in os.listdir('.'):
        l.append(file)

for i in l:
        doc = docx.Document(i)
        name = ''
        la_name = ''
        k = 0
        t = 0
        for j in i:
                if j == chr(32): k += 1
                elif j == chr(46): t += 1
                elif j != chr(32) and k == 0: name += j
                elif j != chr(46) and t == 0: la_name += j

        for i in range(m,p):
              for j in range(1,7):
                    if j == 1:
                        sheet.cell(row=i, column=j).value = name
                    elif j == 2:
                        sheet.cell(row=i, column=j).value = la_name
                    elif j == 3:
                        sheet.cell(row=i, column=j).value = doc.paragraphs[0].text
                    elif j == 4:
                        sheet.cell(row=i, column=j).value = doc.paragraphs[1].text
                    elif j == 5:
                        sheet.cell(row=i, column=j).value = doc.paragraphs[2].text
                    else:
                        sheet.cell(row=i, column=j).value = doc.paragraphs[3].text
        m += 1
        p += 1

os.chdir('C:\Users\Polis\Desktop\Documents\Python\Database - Excel')
wb.save('kod_taxis.xlsx')
